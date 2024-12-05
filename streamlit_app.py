import streamlit as st
from openai import OpenAI
import pandas as pd
from docx import Document
from io import BytesIO
from PIL import Image
import matplotlib.pyplot as plt

# Initialize session state variables safely
st.session_state.setdefault("generated_memo", "")
st.session_state.setdefault("uploaded_files", None)
st.session_state.setdefault("user_changes", "")
st.session_state.setdefault("additional_content", "")

# Function to refresh the page by clearing the session state
def refresh_page():
    # Clear session state variables
    for key in list(st.session_state.keys()):
        del st.session_state[key]

# Access the OpenAI API key securely from Streamlit secrets
client = OpenAI(api_key=st.secrets["openai"]["api_key"])

# Add custom CSS for UI
st.markdown(
    """
    <style>
        body {
            background-color: #E6ECF5;
        }
        .title {
            color: #4E81BD;
            font-size: 32px;
            font-weight: bold;
            text-align: center;
            margin: 10px 0 30px;
        }
    </style>
    """,
    unsafe_allow_html=True,
)

# Use Streamlit columns to center the logo
col1, col2, col3 = st.columns([2, 2, 1])  # Adjust column widths to control spacing
with col2:
    st.image("TalosLogo.png", width=150)  # Centered logo with specific width

# Display title and description
st.markdown("<div class='title'>Talos Labs C&I Co-Pilot</div>", unsafe_allow_html=True)

# File uploader
uploaded_files = st.file_uploader(
    "Please upload relevant documents, including personal financial statements, LP memos, and any client communication regarding requested changes.",
    accept_multiple_files=True,
    type=("txt", "md", "pdf", "xlsx", "docx"),
    key="file_upload",
)

# Function to generate DSCR chart with column normalization
def generate_dscr_chart(df):
    # Normalize column names by stripping whitespace and converting to lowercase
    df.columns = df.columns.str.strip()  # Remove leading/trailing spaces
    df.columns = df.columns.str.lower()  # Convert to lowercase for uniformity
    
    # Check for required columns
    if (
        "year" in df.columns
        and "debt service coverage ratio" in df.columns
        and "minimum debt service coverage ratio" in df.columns
    ):
        try:
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.bar(df["year"], df["debt service coverage ratio"], label="Debt Service Coverage Ratio", alpha=0.7, color="blue")
            ax.plot(
                df["year"], 
                df["minimum debt service coverage ratio"], 
                color="red", 
                marker="o", 
                label="Minimum Debt Service Coverage Ratio", 
                linewidth=2
            )
            ax.set_title("Debt Service Coverage Ratio (DSCR) Over Time", fontsize=14)
            ax.set_xlabel("Year", fontsize=12)
            ax.set_ylabel("Ratio", fontsize=12)
            ax.legend()
            ax.grid(visible=True, linestyle='--', alpha=0.5)
            st.pyplot(fig)
        except Exception as e:
            st.error(f"Error creating chart: {e}")
    else:
        st.error(
            "The uploaded file must contain 'Year', 'Debt Service Coverage Ratio', and 'Minimum Debt Service Coverage Ratio' columns."
        )

# Process uploaded files
data = pd.read_excel("Parse_CSV (3).xlsx", header=3)  # Adjust 'header' index to the correct row
if uploaded_files:
    st.session_state["uploaded_files"] = uploaded_files
    additional_content = ""
    for uploaded_file in uploaded_files:
        st.markdown(f"<div class='small-font'>Filename: {uploaded_file.name}</div>", unsafe_allow_html=True)
        if uploaded_file.type == "text/plain":
            additional_content += uploaded_file.read().decode("utf-8") + "\n"
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                additional_content += para.text + "\n"
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            df = pd.read_excel(uploaded_file)
            additional_content += df.to_string(index=False) + "\n"
            generate_dscr_chart(df)
        else:
            additional_content += "Unsupported file type.\n"
    st.session_state["additional_content"] = additional_content

# Text area for user changes
st.text_area(
    "Please copy and paste change request information.",
    value=st.session_state.get("user_changes", ""),
    placeholder="E.g., include property details, investment summary, changes to net worth, etc.",
    key="user_changes",
)

# Function to handle memo generation
def generate_memo(is_material):
    try:
        memo_type = "Material Change Memo" if is_material else "Non-Material Change Memo"
        document_content = f"Uploaded content:\n{st.session_state['additional_content']}\nUser changes:\n{st.session_state['user_changes'].strip()}"
        messages = [
            {
                "role": "user",
                "content": (
                    f"{document_content}\n\n---\n\n"
                    "Please act as a commercial lender at a top bank. You closed a commercial loan in 2021, and after 3 years, the borrower has requested a 6-month extension to the loan term due to permitting issues that took longer than expected. To maintain a strong relationship with this borrower, you are incentivized to secure credit team approval for the loan term extension. Create a non-material change memo to formalize this 6-month extension."
                    "The memo should have the following structure and include a visual chart to support the case. Ensure formatting is consistent, and headers for all sections are bold."
                    "Structure for the Memo:"
                    "Section 1: Background Information"
                    "Include relevant property information, investment summary, and rationale for the initial investment."
                    "Section 2: Financial Information"
                    "Please include all relevant financial information that we can extract from the memo."
                    "Section 3: Rationale for the Extension"
                    "Outline the borrowers current financial position and why this aligns with the bank's long-term interests."
                ),
            }
        ]
        response = client.chat.completions.create(model="gpt-4", messages=messages)
        st.session_state["generated_memo"] = response.choices[0].message.content
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")

# Buttons for memo generation and refresh
col1, col2, col3 = st.columns(3)
with col1:
    if st.button("Generate Non-Material Change Memo"):
        generate_memo(is_material=False)
with col2:
    if st.button("Generate Material Change Memo"):
        generate_memo(is_material=True)
with col3:
    if st.button("Refresh Page"):
        refresh_page()

# Display generated memo if available
if st.session_state.get("generated_memo", ""):
    st.subheader("Generated Memo")
    st.markdown(f"<div class='left-aligned'>{st.session_state['generated_memo']}</div>", unsafe_allow_html=True)

    # Save the memo as a Word document
    output_doc = Document()
    output_doc.add_heading("Generated Memo", level=1)
    output_doc.add_paragraph(st.session_state["generated_memo"])

    buffer = BytesIO()
    output_doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download Memo",
        data=buffer,
        file_name="Generated_Memo.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )